"""
Загрузка и парсинг документов лотов torgi.gov (PDF, DOC, DOCX).

Документы прикрепляются через массив noticeAttachments в детальном API.
Каждый элемент содержит fileId, fileName, attachmentTypeCode/Name (последние
часто пустые).

Алгоритм классификации (приоритет: код → имя файла):
- notice         — Извещение (описание участка, локация, ВРИ, обременения)
- tech_conditions — ТУ (электричество/газ/вода)
- contract       — Проект договора (переуступка, субаренда, штрафы)
- scheme         — СРЗУ (мало текста, опционально)
"""
import io
import re
from typing import Optional


# Ключевые слова для классификации (искать в любом из: code, type_name, file_name)
TYPE_KEYWORDS = {
    "notice": [
        "notice_document", "notice ", "извещени", "оповещени",
    ],
    "tech_conditions": [
        "technical", "tc_document", "tech_conditions",
        "техническ", "тех условия", "ту ",
    ],
    "contract": [
        "draft_contract", "contract", "договор", "проект дог",
        "договора", "проекта договора", "перенайм",
    ],
    "scheme": [
        "scheme", "срзу", "схема расположен",
    ],
}


def classify_attachment(att: dict) -> Optional[str]:
    """Возвращает наш ключ типа документа (notice/tech_conditions/contract/scheme) или None."""
    code = (att.get("attachmentTypeCode") or "").lower()
    name = (att.get("attachmentTypeName") or "").lower()
    file_name = (att.get("fileName") or "").lower()
    haystack = f"{code} | {name} | {file_name}"

    for our_type, keywords in TYPE_KEYWORDS.items():
        for kw in keywords:
            if kw in haystack:
                return our_type
    return None


SUPPORTED_EXTENSIONS = (".pdf", ".pdf.zip", ".doc", ".docx", ".rtf")


def get_file_ext(file_name: str) -> str:
    """Возвращает расширение файла в нижнем регистре."""
    fname = (file_name or "").lower().strip()
    for ext in SUPPORTED_EXTENSIONS:
        if fname.endswith(ext):
            return ext
    # Может быть без расширения, или необычное расширение
    m = re.search(r"\.([a-z0-9]{1,5})$", fname)
    return f".{m.group(1)}" if m else ""


def select_best_attachments(attachments: list[dict]) -> dict[str, dict]:
    """
    Из массива noticeAttachments выбирает по 1 документу каждого нашего типа.
    Если для типа "contract"/"notice" не нашлось по классификатору — берём
    дополнительно первый файл с подходящим именем.
    """
    result: dict[str, dict] = {}
    if not attachments:
        return result

    # Первый проход — точные совпадения по классификатору
    for att in attachments:
        if att.get("inactive"):
            continue
        ext = get_file_ext(att.get("fileName", ""))
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        our_type = classify_attachment(att)
        if our_type and our_type not in result:
            result[our_type] = att

    return result


async def download_file(client, file_id: str) -> Optional[bytes]:
    """Скачивает файл по fileId через прокси-клиент скрапера."""
    if not file_id:
        return None
    url = f"https://torgi.gov.ru/new/file-store/v1/{file_id}"
    try:
        resp = await client.get(url, timeout=90)
        if resp.status_code != 200:
            return None
        return resp.content
    except Exception:
        return None


# Backward-compat alias
download_pdf = download_file


def _extract_pdf(content: bytes, max_pages: int = 30) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            chunks = []
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    break
                try:
                    text = page.extract_text() or ""
                    chunks.append(text)
                except Exception:
                    pass
            return "\n".join(chunks)
    except Exception as e:
        print(f"[pdf-parse] PDF error: {type(e).__name__}: {e}")
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        chunks = [p.text for p in doc.paragraphs if p.text]
        # Также таблицы
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        chunks.append(cell.text)
        return "\n".join(chunks)
    except Exception as e:
        print(f"[pdf-parse] DOCX error: {type(e).__name__}: {e}")
        return ""


def _is_meaningful_text(text: str, min_ratio: float = 0.55) -> bool:
    """
    Проверяет что текст похож на читаемый русский/латинский, а не на бинарный
    мусор после декодирования OLE/.doc.

    Считаем долю символов из «читаемого» алфавита: кириллица, латиница, цифры,
    стандартная пунктуация, пробел/перенос. Если она ниже порога — отбрасываем.
    """
    if not text or len(text) < 200:
        return False
    sample = text[:5000]
    good = sum(
        1 for ch in sample
        if (
            "а" <= ch.lower() <= "я"
            or "a" <= ch.lower() <= "z"
            or ch.isdigit()
            or ch in " \n\r\t.,;:!?-—–«»\"'()/№%"
            or ch == "ё"
        )
    )
    return good / len(sample) >= min_ratio


def _extract_doc_via_antiword(content: bytes) -> str:
    """Извлечение текста из .doc через утилиту antiword (если установлена)."""
    import subprocess
    try:
        r = subprocess.run(
            ["antiword", "-m", "UTF-8", "-"],
            input=content,
            capture_output=True,
            timeout=30,
        )
        if r.returncode == 0 and r.stdout:
            return r.stdout.decode("utf-8", errors="ignore")
    except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
        if not isinstance(e, FileNotFoundError):
            print(f"[pdf-parse] antiword error: {type(e).__name__}: {e}")
    return ""


def _extract_rtf(content: bytes) -> str:
    """Грубое извлечение видимого текста из .rtf."""
    try:
        text = content.decode("cp1251", errors="ignore")
        # Удаляем control words \xxx и группы заголовка
        text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)
        text = re.sub(r"[{}]", " ", text)
        # Hex escapes \'XX (cp1251)
        def hex_to_char(m):
            try:
                return bytes([int(m.group(1), 16)]).decode("cp1251", errors="ignore")
            except Exception:
                return ""
        text = re.sub(r"\\'([0-9a-fA-F]{2})", hex_to_char, text)
        text = re.sub(r"\s+", " ", text).strip()
        return text
    except Exception as e:
        print(f"[pdf-parse] RTF error: {type(e).__name__}: {e}")
        return ""


def _extract_doc_or_rtf(content: bytes) -> str:
    """
    Парсит старый .doc (OLE compound) и .rtf.

    .rtf начинается с '{\\rtf' — для него простая разборка.
    .doc — OLE compound с сигнатурой D0CF11E0. Для него зовём antiword,
    fallback на байтный скан с валидацией.
    """
    if not content:
        return ""

    # RTF — простой текстовый формат с control words
    if content[:5] in (b"{\\rtf", b"{\\rtf1"):
        text = _extract_rtf(content)
        return text if _is_meaningful_text(text) else ""

    # OLE compound (.doc 97-2003) — D0 CF 11 E0 A1 B1 1A E1
    is_ole = content[:8] == b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"

    if is_ole:
        text = _extract_doc_via_antiword(content)
        if _is_meaningful_text(text):
            return text
        # antiword не справился или не установлен — не сохраняем мусор
        return ""

    # Прочий бинарный мусор — пробуем декод и валидируем
    for encoding in ("cp1251", "utf-8", "latin-1"):
        try:
            text = content.decode(encoding, errors="ignore")
            words = re.findall(r"[А-Яа-яЁёA-Za-z0-9.,%/\-«»()№\s]{4,}", text)
            joined = " ".join(words)
            if _is_meaningful_text(joined):
                return joined
        except Exception:
            continue
    return ""


def extract_text(content: bytes, file_name: str = "", max_pages: int = 30) -> str:
    """
    Извлекает текст из документа любого поддерживаемого формата.
    Определяет формат по расширению + магическим байтам.
    Возвращает "" если извлечь читаемый текст не удалось.
    """
    if not content or len(content) < 100:
        return ""

    ext = get_file_ext(file_name)

    # PDF — по расширению или магическим байтам
    if ext == ".pdf" or content[:4] == b"%PDF":
        text = _extract_pdf(content, max_pages)
        return text if _is_meaningful_text(text) else ""

    # DOCX — это zip с word/document.xml внутри
    if ext == ".docx" or content[:4] == b"PK\x03\x04":
        text = _extract_docx(content)
        return text if _is_meaningful_text(text) else ""

    # DOC/RTF/прочее — внутри уже валидируется
    return _extract_doc_or_rtf(content)


# Backward-compat
def extract_text_from_pdf(pdf_bytes: bytes, max_pages: int = 30) -> str:
    return extract_text(pdf_bytes, ".pdf", max_pages)


def truncate_for_db(text: str, limit: int = 100_000) -> str:
    """Обрезает текст до разумного размера для хранения в БД."""
    if not text:
        return ""
    text = text.strip()
    if len(text) <= limit:
        return text
    return text[: limit - 1000] + "\n\n[...пропущено...]\n\n" + text[-1000:]
